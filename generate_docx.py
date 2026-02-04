from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    doc = Document()

    # Title Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('ARI3215: Robotics Assignment\n')
    run.font.size = Pt(24)
    run.bold = True
    
    run = title_para.add_run('Project: Sort-a-bot\n\n')
    run.font.size = Pt(18)
    run.bold = True

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('Team Members:\n')
    run.font.size = Pt(14)
    run.bold = True
    
    names = [
        "1. Ellyn Rose Debrincat",
        "2. Joachim Grech",
        "3. Benjamin Zammit"
    ]
    for name in names:
        run = para.add_run(f'{name}\n')
        run.font.size = Pt(12)

    para.add_run('\n---\n')
    run = para.add_run('University of Malta | B.Sc. AI | February 2026')
    run.font.italic = True
    
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Sort-a-bot is an autonomous robotic system designed to navigate a dynamic arena, '
        'locate objects (dumbbells), and transport them to designated locations. This project '
        'implements a full ROS 2 integration with a physics-based Gazebo Harmonic simulation.'
    )

    # System Architecture
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph('The system is divided into three primary ROS 2 packages:')
    doc.add_paragraph('sortabot_description: Contains the URDF/Xacro models, inertia parameters, and Gazebo plugins.', style='List Bullet')
    doc.add_paragraph('sortabot_simulation: Manages the Gazebo worlds, launch files, and bridge configurations.', style='List Bullet')
    doc.add_paragraph('sortabot_actions: Contains the logical "brain" of the robot, including PID controllers, action servers, and navigation scripts.', style='List Bullet')

    # Installation & Setup
    doc.add_heading('3. Installation & Setup', level=1)
    doc.add_heading('Prerequisites', level=2)
    doc.add_paragraph('ROS 2 (Jazzy or Harmonic), Gazebo Sim (Harmonic)')
    doc.add_paragraph('Python Dependencies: pip3 install opencv-python PyYAML numpy')
    
    doc.add_heading('Build Instructions', level=2)
    doc.add_paragraph('cd ~/ros2_ws\nrm -rf build/ install/ log/\ncolcon build --symlink-install\nsource install/setup.bash')

    # Simulation Environment (Gazebo)
    doc.add_heading('4. Simulation Environment (Gazebo)', level=1)
    doc.add_paragraph(
        'Our simulation features a balanced, weighted robot model (5kg) with optimized '
        'inertia tensors to ensure stability during high-speed turns.'
    )
    doc.add_paragraph('To launch the simulation: ros2 launch sortabot_simulation sortabot.launch.py')

    # Navigation Methods
    doc.add_heading('5. Navigation Methods', level=1)
    
    doc.add_heading('LIDAR A* Navigator', level=2)
    doc.add_paragraph(
        'A robust pathfinding implementation using an occupancy grid generated from live '
        'LIDAR sweeps and the A* search algorithm for deterministic goal reaching.'
    )
    doc.add_paragraph('Run command:\ncd ~/ros2_ws/src/robotics-assignment/sortabot_actions/scripts\npython3 lidar_navigator.py')

    doc.add_heading('Reinforcement Learning Bridge', level=2)
    doc.add_paragraph(
        'A lightweight PPO-trained inference bridge using a Numpy-only implementation for '
        'high-speed decision making without heavy ML dependencies on the VM.'
    )
    doc.add_paragraph('Run command:\n1. Start Controller: ros2 run sortabot_actions action_server.py\n2. Start RL Brain: python3 drive_robot_numpy.py')

    # Troubleshooting
    doc.add_heading('6. Troubleshooting', level=1)
    doc.add_paragraph('Robot Tipping: The URDF has been weighted to 5kg with a low Center of Mass.', style='List Bullet')
    doc.add_paragraph('Line Endings: All scripts have been standardized to LF for Linux compatibility.', style='List Bullet')
    doc.add_paragraph('Odom Stale Warning: Ensure the ros_gz_bridge is running correctly.', style='List Bullet')

    doc.add_paragraph('\n---\n*Documentation finalized on 2026-02-04.*')

    doc.save('Sort-a-bot_Final_Documentation.docx')
    print("Document saved successfully!")

if __name__ == "__main__":
    create_docx()
